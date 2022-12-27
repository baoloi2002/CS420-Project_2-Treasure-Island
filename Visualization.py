from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, RGBColor
from docx.shared import Inches
import MapPaint
import random

def input():
    global N, M, regionMap, specialMap, pirateReveal, pirateFree, numRegion, Tx, Ty, boundaryMap
    global maskMap
    with open("Map.txt", "r") as f:
        N, M = [int(u) for u in f.readline().split()]
        pirateReveal = int(f.readline()) 
        pirateFree = int(f.readline()) 
        numRegion = int(f.readline())
        Tx, Ty = [int(u) for u in f.readline().split()]
        # MAP with number
        regionMap = [[0 for j in range(M)] for i in range(N)]
        # Map with prison, treasure, mountain
        specialMap = [[' ' for j in range(M)] for i in range(N)]
        # Map with boundary
        boundaryMap = [[0 for j in range(M)] for i in range(N)]
        maskMap = [[0 for j in range(M)] for i in range(N)]

        for i in range(N):
            s = f.readline().replace(" ", "")
            s = s.replace("\n", "")
            tmp = [u for u in s.split(";")]
            for j in range(M):
                if tmp[j][-1] not in "0123456789":
                    specialMap[i][j] = tmp[j][-1]
                    regionMap[i][j] = int(tmp[j][:len(tmp[j])-1])
                else:
                    regionMap[i][j] = int(tmp[j])


def clearBoundary():
    global boundaryMap
    for i in range(N):
        for j in range(M):
            boundaryMap[i][j] = 0

def hint_1():
    global LOG, boundaryMap, document
    sLog = LOG.pop(0)
    document.add_paragraph(sLog)
    sz = int(sLog)
    lst  = []
    clearBoundary()
    for _ in range(sz):
        sLog = LOG.pop(0)
        document.add_paragraph(sLog)
        u = [int(x) for x in sLog.strip().split(' ')]
        boundaryMap[u[0]][u[1]] = 1
    
    MapPaint.main(1600, 1600, N, M, regionMap, specialMap, boundaryMap, maskMap, Tx, Ty, Ax, Ay, colorMap)
    document.add_picture("test.png", width=Inches(7))

def hint_2():
    global LOG, boundaryMap, document
    sLog = LOG.pop(0)
    document.add_paragraph(sLog)
    sz = int(sLog)
    lst  = []
    clearBoundary()
    sLog = LOG.pop(0)
    document.add_paragraph(sLog)
    u = [int(x) for x in sLog.strip().split(' ')]
    for i in range(N):
        for j in range(M):
            if regionMap[i][j] in u:
                boundaryMap[i][j] = 1
    
    MapPaint.main(1600, 1600, N, M, regionMap, specialMap, boundaryMap, maskMap, Tx, Ty, Ax, Ay, colorMap)
    document.add_picture("test.png", width=Inches(7))

def hint_3():
    global LOG, boundaryMap, document
    sLog = LOG.pop(0)
    document.add_paragraph(sLog)
    sz = int(sLog)
    lst  = []
    clearBoundary()
    sLog = LOG.pop(0)
    document.add_paragraph(sLog)
    u = [int(x) for x in sLog.strip().split(' ')]
    for i in range(N):
        for j in range(M):
            if regionMap[i][j] in u:
                boundaryMap[i][j] = 1
    
    MapPaint.main(1600, 1600, N, M, regionMap, specialMap, boundaryMap, maskMap, Tx, Ty, Ax, Ay, colorMap)
    document.add_picture("test.png", width=Inches(7))

def hint_4():
    global LOG, boundaryMap, document
    sLog = LOG.pop(0)
    document.add_paragraph(sLog)
    u = [int(x) for x in sLog.strip().split(' ')]
    sLog = LOG.pop(0)
    document.add_paragraph(sLog)
    v = [int(x) for x in sLog.strip().split(' ')]

    if u[0] > v[0]:
        u[0], v[0] = v[0], u[0]
    if u[1] > v[1]:
        u[1], v[1] = v[1], u[1]

    clearBoundary()
    for i in range(N):
        if u[0] <= i and i <= v[0]:
            for j in range(M):
                if u[1] <= j and j <= v[1]:
                    boundaryMap[i][j] = 1
    
    MapPaint.main(1600, 1600, N, M, regionMap, specialMap, boundaryMap, maskMap, Tx, Ty, Ax, Ay, colorMap)
    document.add_picture("test.png", width=Inches(7))

def hint_5():
    global LOG, boundaryMap, document
    sLog = LOG.pop(0)
    document.add_paragraph(sLog)
    u = [int(x) for x in sLog.strip().split(' ')]
    sLog = LOG.pop(0)
    document.add_paragraph(sLog)
    v = [int(x) for x in sLog.strip().split(' ')]

    if u[0] > v[0]:
        u[0], v[0] = v[0], u[0]
    if u[1] > v[1]:
        u[1], v[1] = v[1], u[1]

    clearBoundary()
    for i in range(N):
        if u[0] <= i and i <= v[0]:
            for j in range(M):
                if u[1] <= j and j <= v[1]:
                    boundaryMap[i][j] = 1
    
    MapPaint.main(1600, 1600, N, M, regionMap, specialMap, boundaryMap, maskMap, Tx, Ty, Ax, Ay, colorMap)
    document.add_picture("test.png", width=Inches(7))

def hint_6():
    global LOG, boundaryMap, document
    sLog = LOG.pop(0)
    document.add_paragraph(sLog)

def hint_7():
    global LOG, boundaryMap, document
    clearBoundary()
    sLog = LOG.pop(0)
    document.add_paragraph(sLog)
    if sLog == 'ROW':
        sLog = LOG.pop(0)
        document.add_paragraph(sLog)
        u = int(sLog)
        for i in range(M):
            boundaryMap[u][i] = 1
    else:
        sLog = LOG.pop(0)
        document.add_paragraph(sLog)
        u = int(sLog)
        for i in range(N):
            boundaryMap[i][u] = 1
            
    MapPaint.main(1600, 1600, N, M, regionMap, specialMap, boundaryMap, maskMap, Tx, Ty, Ax, Ay, colorMap)
    document.add_picture("test.png", width=Inches(7))

def hint_8():
    global LOG, boundaryMap, document
    clearBoundary()
    sLog = LOG.pop(0)
    document.add_paragraph(sLog)
    if sLog == 'ROW':
        sLog = LOG.pop(0)
        document.add_paragraph(sLog)
        u = int(sLog)
        for i in range(M):
            boundaryMap[u][i] = 1
    else:
        sLog = LOG.pop(0)
        document.add_paragraph(sLog)
        u = int(sLog)
        for i in range(N):
            boundaryMap[i][u] = 1

    MapPaint.main(1600, 1600, N, M, regionMap, specialMap, boundaryMap, maskMap, Tx, Ty, Ax, Ay, colorMap)
    document.add_picture("test.png", width=Inches(7))

def hint_9():
    global LOG, boundaryMap, document
    clearBoundary()
    sLog = LOG.pop(0)
    document.add_paragraph(sLog)
    u = [int(x) for x in sLog.strip().split(' ')]
    for i in range(N):
        for j in range(M):
            isK = False
            for x in range(i-1, i+2):
                for y in range(j-1, j+2):
                    if x >= 0 and x < N and y >= 0 and y < M:
                        if regionMap[x][y] in u:
                            isK = True
            if isK:
                boundaryMap[i][j] = True

    MapPaint.main(1600, 1600, N, M, regionMap, specialMap, boundaryMap, maskMap, Tx, Ty, Ax, Ay, colorMap)
    document.add_picture("test.png", width=Inches(7))

def hint_10():
    global LOG
    pass

def hint_11():
    global LOG
    pass

def hint_12():
    global LOG, boundaryMap, document
    sLog = LOG.pop(0)
    document.add_paragraph(sLog)
    u = [int(x) for x in sLog.strip().split(' ')]
    sLog = LOG.pop(0)
    document.add_paragraph(sLog)
    v = [int(x) for x in sLog.strip().split(' ')]

    if u[0] > v[0]:
        u[0], v[0] = v[0], u[0]
    if u[1] > v[1]:
        u[1], v[1] = v[1], u[1]

    clearBoundary()
    for i in range(N):
        if u[0] <= i and i <= v[0]:
            for j in range(M):
                if u[1] <= j and j <= v[1]:
                    boundaryMap[i][j] = 1
    
    MapPaint.main(1600, 1600, N, M, regionMap, specialMap, boundaryMap, maskMap, Tx, Ty, Ax, Ay, colorMap)
    document.add_picture("test.png", width=Inches(7))

def hint_13():
    global LOG, boundaryMap, document
    sLog = LOG.pop(0)
    document.add_paragraph(sLog)
    u = [int(x) for x in sLog.strip().split(' ')]
    sLog = LOG.pop(0)
    document.add_paragraph(sLog)
    direct = sLog
    clearBoundary()
    for i in range(N):
        for j in range(M):
            X = i - u[0]
            Y = j - u[1]
            lst = []
            if X >= 0:
                if Y >= 0:
                    lst.append('SE')
                    if Y <= X:
                        lst.append('S')
                    if Y >= X:
                        lst.append('E')

                if Y <= 0:
                    lst.append('SW')
                    if -Y <= X:
                        lst.append('S')
                    if -Y >= X:
                        lst.append('W')
            if X <= 0:
                if Y >= 0:
                    lst.append('NE')
                    if Y <= -X:
                        lst.append('N')
                    if Y >= -X:
                        lst.append('E')
                if Y <= 0:
                    lst.append('NW')
                    if -Y <= -X:
                        lst.append('N')
                    if -Y >= -X:
                        lst.append('W')
            
            if direct in lst:
                boundaryMap[i][j] = 1
    
    MapPaint.main(1600, 1600, N, M, regionMap, specialMap, boundaryMap, maskMap, Tx, Ty, Ax, Ay, colorMap)
    document.add_picture("test.png", width=Inches(7))


def hint_14():
    global LOG, boundaryMap, document
    sLog = LOG.pop(0)
    document.add_paragraph(sLog)
    uX = [int(x) for x in sLog.strip().split(' ')]
    sLog = LOG.pop(0)
    document.add_paragraph(sLog)
    vX = [int(x) for x in sLog.strip().split(' ')]
    sLog = LOG.pop(0)
    document.add_paragraph(sLog)
    u = [int(x) for x in sLog.strip().split(' ')]
    sLog = LOG.pop(0)
    document.add_paragraph(sLog)
    v = [int(x) for x in sLog.strip().split(' ')]

    clearBoundary()
    for i in range(N):
            for j in range(M):
                if ((uX[0] <= i and i < u[0]) or (v[0] < i and i <= vX[0])) and uX[1] <= j and j <= vX[1]:
                    boundaryMap[i][j] = 1
                if ((uX[1] <= j and j < u[1]) or (v[1] < j and j <= vX[1])) and uX[0] <= i and i <= vX[0]:
                    boundaryMap[i][j] = 1
    
    MapPaint.main(1600, 1600, N, M, regionMap, specialMap, boundaryMap, maskMap, Tx, Ty, Ax, Ay, colorMap)
    document.add_picture("test.png", width=Inches(7))



def hint_15():
    global LOG, boundaryMap, document
    sLog = LOG.pop(0)
    document.add_paragraph(sLog)
    u = int(sLog)
    clearBoundary()
    for i in range(N):
        for j in range(M):
            if regionMap[i][j] == u:
                boundaryMap[i][j] =1 
    
    MapPaint.main(1600, 1600, N, M, regionMap, specialMap, boundaryMap, maskMap, Tx, Ty, Ax, Ay, colorMap)
    document.add_picture("test.png", width=Inches(7))


def hintVisual():
    global LOG, document
    sLog = LOG.pop(0)
    document.add_paragraph(sLog)
    v = int(sLog)
    if v == 1:
        hint_1()
    if v == 2:
        hint_2()
    if v == 3:
        hint_3()
    if v == 4:
        hint_4()
    if v == 5:
        hint_5()
    if v == 6:
        hint_6()
    if v == 7:
        hint_7()
    if v == 8:
        hint_8()
    if v == 9:
        hint_9()
    if v == 10:
        hint_10()
    if v == 11:
        hint_11()
    if v == 12:
        hint_12()
    if v == 13:
        hint_13()
    if v == 14:
        hint_14()
    if v == 15:
        hint_15()

def maskPaint():
    global LOG, boundaryMap, document, Ax, Ay
    sLog = LOG.pop(0)
    document.add_paragraph(sLog)
    u = [int(x) for x in sLog.strip().split(" ")]
    Ax = u[0]
    Ay = u[1]
    for i in range(N):
        sLog = LOG.pop(0)
        #document.add_paragraph(sLog)
        a = [int(u) for u in sLog.strip().split(" ")]
        for j in range(M):
            if a[j]:
                maskMap[i][j] = 1
    
    MapPaint.main(1600, 1600, N, M, regionMap, specialMap, boundaryMap, maskMap, Tx, Ty, Ax, Ay, colorMap)
    document.add_picture("test.png", width=Inches(7))

def main(log):
    global document, LOG, colorMap, Ax, Ay

    
    LOG = log
    input()

    colorMap = [(255, 204, 204), (153, 255, 51), (0, 255, 0), (255, 204, 153), (255, 255, 102),
                (0, 204, 102), (0, 153, 153), (255, 51, 51), (192, 192, 192), (96, 96, 96),
                (255, 0, 127), (255, 204, 255), (102, 102, 255), (51, 255, 153), (51, 51, 0),
                (255, 230, 102), (12, 26, 22), (7, 57, 20), (64, 83, 13), (44, 43, 76)]
    random.shuffle(colorMap)
    colorMap = [(0, 0, 102)] + colorMap
    
    document = Document()
    
    while len(LOG) > 0:
        u = LOG.pop(0)
        document.add_paragraph(u)
        if "AGENT:" in u:
            tmp = [v for v in u.strip().split()]
            Ax = int(tmp[1])
            Ay = int(tmp[2])
        if "HINT" in u:
            hintVisual()
        if "MASK" in u:
            maskPaint()


        #MapPaint.main(1600, 1600, N, M, regionMap, specialMap, boundaryMap, Tx, Ty, colorMap)
        #document.add_picture("test.png", width=Inches(7))

        #document.add_page_break()




    document.save("Visualization.docx")
